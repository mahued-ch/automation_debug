from docx import Document
from docx.shared import Inches
from Automation_Variables import ruta, ruta_output
import os
import re

def extraer_numero(nombre_archivo):
    """
    Extrae el número de un nombre de archivo.
    Supone que el número está al final del nombre, antes de la extensión.
    """
    # Utiliza una expresión regular para extraer el número al final del nombre
    match = re.search(r'(\d+)(?=\.[a-zA-Z]+$)', nombre_archivo)
    return int(match.group(1)) if match else float('inf')  # Retorna un valor muy alto si no se encuentra número

def Genera_Documento(nombre, prefijo, cabecera):
    doc = Document()
    doc.add_heading(cabecera, level=1)
# Obtener la lista de archivos de imágenes en la ruta
#    imagenes = [img for img in os.listdir(ruta_output) if img.lower().endswith(('.png', '.jpg', '.jpeg'))]
    imagenes = [img for img in os.listdir(ruta_output) 
                if img.lower().startswith(prefijo.lower()) and img.lower().endswith(('.png', '.jpg', '.jpeg'))]
    
    # Ordenar las imágenes de manera inteligente por el número al final del nombre
    imagenes.sort(key=extraer_numero)    
    
# Insertar cada imagen en el documento
    for imagen in imagenes:
        doc.add_paragraph(f'Captura: {imagen}')
        doc.add_picture(os.path.join(ruta_output, imagen), width=Inches(6))  # Ajusta el ancho según sea necesario
# Guardar el archivo .docx
    doc.save(os.path.join(ruta_output, nombre))
    print(f'Documento generado: {os.path.join(ruta_output, nombre)}')
